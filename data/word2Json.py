# -*- coding: utf-8 -*-            
# @Author : Sofiya
# @Time : 2024/8/7 上午9:03
# Note:训练数据准备

import os
from docx import Document
import json

def word_to_json(word_file_path):
    """
    将单个Word文档转换为JSON数据。
    """
    doc = Document(word_file_path)
    json_data = {}

    # 提取文档中的所有段落
    for i, paragraph in enumerate(doc.paragraphs):
        json_data[f'paragraph_{i}'] = paragraph.text

    # 提取文档中的所有表格
    for j, table in enumerate(doc.tables):
        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            json_data[f'table_{j}_{len(row_data)}'] = row_data

    return json_data

def convert_folder_to_json(folder_path, output_json_file):
    """
    将文件夹中所有Word文档转换为JSON，并保存到一个文件。
    """
    # 用于存储所有文档的JSON数据
    combined_json_data = {}

    # 遍历文件夹中的所有文件
    for filename in os.listdir(folder_path):
        if filename.endswith('.docx'):
            word_file_path = os.path.join(folder_path, filename)
            # 将当前文档转换为JSON
            doc_json_data = word_to_json(word_file_path)
            # 将转换得到的数据添加到总的JSON数据中
            combined_json_data[filename] = doc_json_data

    # 将所有文档的JSON数据转换为JSON格式的字符串
    json_string = json.dumps(combined_json_data, ensure_ascii=False, indent=4)

    # 保存JSON字符串到文件
    with open(output_json_file, 'w', encoding='utf-8') as json_file:
        json_file.write(json_string)
    print(f"All Word documents have been converted and saved to '{output_json_file}'.")



if __name__ == '__main__':

    # 指定文件夹路径和输出JSON文件名
    folder_path = '/Users/apple/Downloads/data'
    output_json_file = '/Users/apple/Downloads/data/test1.json'

    # 执行转换
    convert_folder_to_json(folder_path, output_json_file)



